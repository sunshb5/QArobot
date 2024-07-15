import os
import docx

local_path = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))


def _docx2text(file_path=None):
    """
    :param file_path: 文件路径
    :return:
    """
    document = docx.Document(file_path)

    data_str = []
    for i in range(len(document.paragraphs)):
        if document.paragraphs[i].text != '':
            text_one = document.paragraphs[i].text
            data_str.append(text_one)

    return data_str


def load_data():
    texts = []
    folder_names = ['中大简介']
    for folder_name in folder_names:
        folder_path = local_path + '/data/%s' % folder_name
        file_paths = os.listdir(folder_path)
        for file_path in file_paths:
            try:
                text = _docx2text(file_path=folder_path + '/' + file_path)
                texts += text
            except (FileNotFoundError, PermissionError, IOError) as e:
                print(f'错误：{folder_path}/{file_path} - {e}')
                texts.append('')
    return texts
